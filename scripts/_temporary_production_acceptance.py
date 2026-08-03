#!/usr/bin/env python3
"""Temporary, self-contained authenticated production acceptance runner.

The runner creates a disposable Supabase user through the public Auth API,
waits for an operator-controlled SQL confirmation/credit grant, and then sends
fictional fixtures through the real production analyzer endpoint. It never
prints the password, access token, email address, document text, filenames, or
customer data. The branch and disposable account are removed after acceptance.
"""

from __future__ import annotations

import io
import json
import os
import secrets
import sys
import time
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

SUPABASE_URL = "https://fqwkvyypjnxkiojbubdf.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImZxd2t2eXlwam54a2lvamJ1YmRmIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODAxOTA2OTIsImV4cCI6MjA5NTc2NjY5Mn0.Pe9DyxonYduw427m7H7GntD40XPS0G6oa3wcBsJ79FY"
PRODUCTION_URL = "https://www.subshield.net"
EXPECTED_STARTING_CREDITS = 8

CANONICAL_TEXT = """SUBSHIELD PRODUCTION ACCEPTANCE FIXTURE - CYBER RISK PACKAGE

Fictional Test Document - For SubShield Production Acceptance Only
NOT A REAL CONTRACT - DO NOT USE FOR BUSINESS OR LEGAL PURPOSES

DOCUMENT CONTROL
Fixture ID: PROD-ACCEPT-2026-08-03
Subcontract Number: QA-2026-803
Prime Contract Number: FA9999-26-C-0803
Prime Contractor: Northstar Civic Systems, Inc. (fictional)
Subcontractor: Redwood Mission Analytics, LLC (fictional)
Contract Type: Firm-Fixed-Price
Total Firm-Fixed Price: $420,000

1. INCORPORATION AND DEFERRED DOCUMENTS

1.3 The System Security Plan, CUI marking guide, network boundary diagram, data-flow map, and Prime cyber procedures are not attached at execution. Prime will provide or revise those materials after award as mission needs evolve.

2. REQUIRED CYBER FRAMEWORKS

2.1 Subcontractor shall comply with DFARS 252.204-7012, Safeguarding Covered Defense Information and Cyber Incident Reporting, and shall implement the security requirements of NIST SP 800-171 on every covered contractor information system.

2.4 Prime may add revised cybersecurity frameworks, agency directives, CMMC requirements, cloud-security controls, or customer procedures by email or portal posting. Each added requirement becomes binding upon notice without a price or schedule adjustment.

4. CYBER INCIDENT REPORTING AND RESPONSE

4.1 Subcontractor shall report any suspected cyber incident, compromise, unauthorized disclosure, malware event, lost device, anomalous access, or policy violation to Prime within eight hours after discovery or suspicion.

4.4 Prime may direct containment, isolation, credential reset, system shutdown, evidence collection, employee interviews, forensic imaging, customer notification, or restoration actions. Subcontractor shall comply immediately and bear all associated cost.

5. CYBERSECURITY ASSESSMENTS, ACCESS, AND REMEDIATION

5.1 Prime, the Government, and their designees may conduct announced or unannounced cybersecurity assessments of Subcontractor systems, facilities, personnel, policies, and lower-tier suppliers at any time.

5.2 Subcontractor shall provide administrative access, network diagrams, System Security Plans, plans of action and milestones, vulnerability scans, penetration-test results, security logs, training records, incident records, and other requested evidence without additional charge.

5.3 If Prime determines that Subcontractor's cybersecurity posture is deficient, Prime may issue a remediation directive or suspend affected system access.

5.4 A suspension or remediation directive does not excuse schedule performance and does not entitle Subcontractor to an equitable adjustment.

8. PAYMENT WITHHOLDING AND CONTINUED PERFORMANCE

8.1 Prime may withhold up to twenty percent of any invoice until Subcontractor provides all cyber evidence requested for the billing period.

8.2 Prime may withhold all payment during an incident investigation, an assessment dispute, a score deficiency, or a suspected failure to flow down requirements.

8.3 Withholding does not relieve Subcontractor from continued performance, incident response, remediation, or reporting obligations.

9. CYBER INDEMNITY AND COST ALLOCATION

9.1 Subcontractor shall indemnify, defend, and hold harmless Prime, the Government, and their personnel from all claims, losses, response costs, notification costs, credit monitoring, forensic costs, restoration expenses, penalties, damages, and attorneys' fees arising from an actual or suspected cyber incident involving Subcontractor or a lower-tier supplier.

9.3 No limitation of liability applies to cybersecurity, confidentiality, data handling, incident reporting, intellectual property, or indemnity obligations.

10. TERMINATION AND CONTINUED PERFORMANCE

10.1 Prime may terminate immediately, without any right to cure, if Prime believes Subcontractor presents an unacceptable cyber risk, reports an incident late, fails an assessment, or does not provide requested access.

10.2 Prime may terminate for convenience on five calendar days notice and is not liable for unabsorbed overhead, demobilization, security investments, unused licenses, or profit on unperformed work.

10.3 Pending any dispute or investigation, Subcontractor shall continue performance, finance all required remediation, and follow Prime direction without delaying mission work.

END OF FICTIONAL PRODUCTION ACCEPTANCE FIXTURE
"""


@dataclass(frozen=True)
class Fixture:
    key: str
    filename: str
    content_type: str
    content: bytes
    pasted: bool = False


def normalize(value: str) -> str:
    return " ".join(value.replace("’", "'").replace("“", '"').replace("”", '"').split()).lower()


def wrap_line(text: str, font: str, size: float, max_width: float) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if stringWidth(candidate, font, size) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def make_docx() -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(9)
    for block in CANONICAL_TEXT.strip().split("\n\n"):
        lines = block.splitlines()
        paragraph = document.add_paragraph()
        heading = len(lines) == 1 and (
            lines[0].isupper() or (len(lines[0]) > 2 and lines[0][0].isdigit() and ". " in lines[0])
        )
        for index, line in enumerate(lines):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            if heading:
                run.bold = True
                run.font.size = Pt(10)
        paragraph.paragraph_format.space_after = Pt(4)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_text_pdf() -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    x = 0.65 * inch
    y = height - 0.6 * inch
    for raw_line in CANONICAL_TEXT.splitlines():
        if not raw_line:
            y -= 5
            continue
        heading = raw_line.isupper() or (
            len(raw_line) > 2 and raw_line[0].isdigit() and ". " in raw_line
        )
        font = "Helvetica-Bold" if heading else "Helvetica"
        size = 9 if heading else 8.2
        for line in wrap_line(raw_line, font, size, width - 1.3 * inch):
            if y < 0.55 * inch:
                pdf.showPage()
                y = height - 0.6 * inch
            pdf.setFont(font, size)
            pdf.drawString(x, y, line)
            y -= 10
    pdf.save()
    return buffer.getvalue()


def make_scanned_pdf() -> bytes:
    regular = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 19)
    bold = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 20)
    page_width, page_height = 1200, 1550
    margin = 70
    max_width = page_width - (2 * margin)
    pages: list[Image.Image] = []
    image = Image.new("1", (page_width, page_height), 1)
    draw = ImageDraw.Draw(image)
    y = 55
    for raw_line in CANONICAL_TEXT.splitlines():
        if not raw_line:
            y += 8
            continue
        heading = raw_line.isupper() or (
            len(raw_line) > 2 and raw_line[0].isdigit() and ". " in raw_line
        )
        font = bold if heading else regular
        words = raw_line.split()
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        for line in lines:
            if y > page_height - 65:
                pages.append(image)
                image = Image.new("1", (page_width, page_height), 1)
                draw = ImageDraw.Draw(image)
                y = 55
            draw.text((margin, y), line, font=font, fill=0)
            y += 29 if heading else 27
    pages.append(image)

    png_buffers: list[io.BytesIO] = []
    for page in pages:
        page_buffer = io.BytesIO()
        page.save(page_buffer, format="PNG", optimize=True)
        page_buffer.seek(0)
        png_buffers.append(page_buffer)

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=letter, pageCompression=1)
    from reportlab.lib.utils import ImageReader
    for page_buffer in png_buffers:
        pdf.drawImage(ImageReader(page_buffer), 0, 0, width=letter[0], height=letter[1])
        pdf.showPage()
    pdf.save()
    return output.getvalue()


def build_fixtures() -> list[Fixture]:
    return [
        Fixture("paste", "Pasted contract text", "application/json", CANONICAL_TEXT.encode(), pasted=True),
        Fixture("txt", "acceptance-fixture.txt", "text/plain", CANONICAL_TEXT.encode()),
        Fixture(
            "docx",
            "acceptance-fixture.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            make_docx(),
        ),
        Fixture("text_pdf", "acceptance-fixture-text.pdf", "application/pdf", make_text_pdf()),
        Fixture("scanned_pdf", "acceptance-fixture-scanned.pdf", "application/pdf", make_scanned_pdf()),
        Fixture("garbled", "acceptance-fixture-garbled.txt", "text/plain", b"\x00\x01\x02UNREADABLE\x00\xff\xfeA B C\n"),
    ]


def request_json(method: str, url: str, **kwargs: Any) -> tuple[int, dict[str, Any]]:
    response = requests.request(method, url, timeout=110, **kwargs)
    try:
        payload = response.json()
    except ValueError:
        payload = {"nonJsonResponse": True}
    return response.status_code, payload if isinstance(payload, dict) else {"value": payload}


def auth_headers(token: str | None = None) -> dict[str, str]:
    headers = {"apikey": SUPABASE_KEY, "Content-Type": "application/json"}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers


def create_disposable_user(run_id: str) -> tuple[str, str]:
    email = f"subshield.acceptance.{run_id}@example.com"
    password = secrets.token_urlsafe(32) + "Aa1!"
    status, payload = request_json(
        "POST",
        f"{SUPABASE_URL}/auth/v1/signup",
        headers=auth_headers(),
        json={
            "email": email,
            "password": password,
            "data": {"purpose": "subshield-production-acceptance"},
        },
    )
    if status not in (200, 201):
        raise RuntimeError(f"Disposable signup failed with HTTP {status}.")
    print("Disposable acceptance signup created; waiting for controlled confirmation and credit grant.")
    return email, password


def wait_for_session_and_credits(email: str, password: str) -> tuple[str, str, int]:
    deadline = time.time() + 900
    last_status = None
    while time.time() < deadline:
        status, payload = request_json(
            "POST",
            f"{SUPABASE_URL}/auth/v1/token?grant_type=password",
            headers=auth_headers(),
            json={"email": email, "password": password},
        )
        last_status = status
        token = payload.get("access_token")
        user = payload.get("user") if isinstance(payload.get("user"), dict) else {}
        user_id = user.get("id")
        if status == 200 and isinstance(token, str) and isinstance(user_id, str):
            claim_status, claim_payload = request_json(
                "POST",
                f"{PRODUCTION_URL}/api/auth/claim",
                headers={"Authorization": f"Bearer {token}"},
            )
            credits = claim_payload.get("credits")
            if claim_status == 200 and isinstance(credits, int) and credits >= EXPECTED_STARTING_CREDITS:
                print("Disposable account confirmed and acceptance credits available.")
                return token, user_id, credits
        time.sleep(10)
    raise RuntimeError(f"Timed out waiting for controlled account preparation; last auth HTTP {last_status}.")


def create_audit(token: str, user_id: str, fixture: Fixture) -> str:
    status, payload = request_json(
        "POST",
        f"{SUPABASE_URL}/rest/v1/contract_audits?select=id",
        headers={
            **auth_headers(token),
            "Prefer": "return=representation",
        },
        json={
            "user_id": user_id,
            "file_name": fixture.filename,
            "status": "Processing",
        },
    )
    rows = payload.get("value") if isinstance(payload.get("value"), list) else payload
    if status not in (200, 201) or not isinstance(rows, list) or not rows:
        raise RuntimeError(f"Audit creation failed for {fixture.key} with HTTP {status}.")
    audit_id = rows[0].get("id")
    if not isinstance(audit_id, (str, int)):
        raise RuntimeError(f"Audit creation returned no safe ID for {fixture.key}.")
    return str(audit_id)


def current_credits(token: str) -> int | None:
    status, payload = request_json(
        "POST",
        f"{PRODUCTION_URL}/api/auth/claim",
        headers={"Authorization": f"Bearer {token}"},
    )
    credits = payload.get("credits")
    return credits if status == 200 and isinstance(credits, int) else None


def quote_grounding_score(quote: str) -> float:
    source = normalize(CANONICAL_TEXT)
    candidate = normalize(quote)
    if not candidate:
        return 0.0
    if candidate in source:
        return 1.0
    source_windows = [
        source[index : index + len(candidate)]
        for index in range(0, max(1, len(source) - len(candidate) + 1), max(1, len(candidate) // 8))
    ]
    return max((SequenceMatcher(None, candidate, window).ratio() for window in source_windows), default=0.0)


def summarize_result(fixture: Fixture, status: int, payload: dict[str, Any], elapsed: float, before: int | None, after: int | None) -> dict[str, Any]:
    wrapped = payload.get("result") if isinstance(payload.get("result"), dict) else {}
    result = wrapped.get("result") if isinstance(wrapped.get("result"), dict) else wrapped
    primary = result.get("primaryTraps") if isinstance(result.get("primaryTraps"), list) else []
    secondary = result.get("secondaryConcerns") if isinstance(result.get("secondaryConcerns"), list) else []
    findings = [item for item in [*primary, *secondary] if isinstance(item, dict)]
    labels = sorted({str(item.get("regulation")) for item in findings if item.get("regulation")})
    quote_scores = [
        quote_grounding_score(str(item.get("foundText", "")))
        for item in findings
        if item.get("foundText")
    ]
    return {
        "case": fixture.key,
        "httpStatus": status,
        "elapsedSeconds": round(elapsed, 2),
        "replayed": bool(wrapped.get("replayed")),
        "creditsBefore": before,
        "creditsAfter": after,
        "creditDelta": None if before is None or after is None else after - before,
        "riskLevel": result.get("riskLevel"),
        "limitedScan": bool(result.get("limitedScan")),
        "limitedScanReasonPresent": bool(result.get("limitedScanReason")),
        "partialOcrScan": bool(result.get("partialOcrScan")),
        "ocrPagesProcessed": result.get("ocrPagesProcessed"),
        "ocrTotalPages": result.get("ocrTotalPages"),
        "findingCount": len(findings),
        "labels": labels,
        "allQuotesNearCanonical": bool(quote_scores) and min(quote_scores) >= 0.90,
        "minimumQuoteGroundingScore": round(min(quote_scores), 3) if quote_scores else None,
        "errorCode": payload.get("code"),
        "creditRestored": payload.get("creditRestored"),
    }


def analyze_fixture(token: str, user_id: str, fixture: Fixture) -> tuple[str, dict[str, Any]]:
    audit_id = create_audit(token, user_id, fixture)
    before = current_credits(token)
    started = time.monotonic()
    if fixture.pasted:
        status, payload = request_json(
            "POST",
            f"{PRODUCTION_URL}/api/analyze-contract",
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"text": CANONICAL_TEXT, "fileName": fixture.filename, "auditId": audit_id},
        )
    else:
        response = requests.post(
            f"{PRODUCTION_URL}/api/analyze-contract",
            headers={"Authorization": f"Bearer {token}"},
            files={
                "file": (fixture.filename, fixture.content, fixture.content_type),
                "auditId": (None, audit_id),
            },
            timeout=110,
        )
        status = response.status_code
        try:
            payload = response.json()
        except ValueError:
            payload = {"nonJsonResponse": True}
    elapsed = time.monotonic() - started
    after = current_credits(token)
    return audit_id, summarize_result(fixture, status, payload, elapsed, before, after)


def replay_completed_case(token: str, fixture: Fixture, audit_id: str) -> dict[str, Any]:
    before = current_credits(token)
    started = time.monotonic()
    status, payload = request_json(
        "POST",
        f"{PRODUCTION_URL}/api/analyze-contract",
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
        json={"text": CANONICAL_TEXT, "fileName": fixture.filename, "auditId": audit_id},
    )
    after = current_credits(token)
    summary = summarize_result(fixture, status, payload, time.monotonic() - started, before, after)
    summary["case"] = "duplicate_replay"
    return summary


def evaluate_consistency(results: list[dict[str, Any]]) -> dict[str, Any]:
    substantive = [item for item in results if item["case"] in {"paste", "txt", "docx", "text_pdf", "scanned_pdf"}]
    baseline = set(next(item for item in substantive if item["case"] == "paste")["labels"])
    comparisons: dict[str, Any] = {}
    for item in substantive:
        labels = set(item["labels"])
        union = baseline | labels
        score = 1.0 if not union else len(baseline & labels) / len(union)
        comparisons[item["case"]] = {
            "jaccardToPaste": round(score, 3),
            "missingFromPasteBaseline": sorted(baseline - labels),
            "additionalVsPasteBaseline": sorted(labels - baseline),
        }
    return comparisons


def write_summary(output: dict[str, Any]) -> None:
    result_path = Path("acceptance-result.json")
    result_path.write_text(json.dumps(output, indent=2, sort_keys=True) + "\n")
    print(json.dumps(output, indent=2, sort_keys=True))

    summary_path = os.environ.get("GITHUB_STEP_SUMMARY")
    if not summary_path:
        return
    lines = [
        "# Authenticated Production Acceptance",
        "",
        f"Starting credits: {output['startingCredits']}",
        f"Ending credits: {output['endingCredits']}",
        "",
        "| Case | HTTP | Credit delta | Findings | Limited | Partial OCR | Seconds |",
        "|---|---:|---:|---:|---|---|---:|",
    ]
    for item in output["results"]:
        lines.append(
            f"| {item['case']} | {item['httpStatus']} | {item['creditDelta']} | "
            f"{item['findingCount']} | {item['limitedScan']} | {item['partialOcrScan']} | "
            f"{item['elapsedSeconds']} |"
        )
    lines.extend(["", "## Label consistency", "", "```json", json.dumps(output["consistency"], indent=2), "```", ""])
    with open(summary_path, "a", encoding="utf-8") as handle:
        handle.write("\n".join(lines))


def main() -> int:
    run_id = os.environ.get("GITHUB_RUN_ID", "local")
    email, password = create_disposable_user(run_id)
    token, user_id, starting_credits = wait_for_session_and_credits(email, password)

    fixtures = build_fixtures()
    results: list[dict[str, Any]] = []
    paste_audit_id: str | None = None
    for fixture in fixtures:
        audit_id, summary = analyze_fixture(token, user_id, fixture)
        results.append(summary)
        if fixture.key == "paste":
            paste_audit_id = audit_id
            results.append(replay_completed_case(token, fixture, audit_id))

    output = {
        "runId": run_id,
        "startingCredits": starting_credits,
        "endingCredits": current_credits(token),
        "results": results,
        "consistency": evaluate_consistency(results),
        "acceptanceChecks": {
            "textRepresentationsExact": all(
                next(item for item in results if item["case"] == key)["httpStatus"] == 200
                and next(item for item in results if item["case"] == key)["allQuotesNearCanonical"]
                for key in ("paste", "txt", "docx", "text_pdf")
            ),
            "scannedCompletedOrLimited": (
                next(item for item in results if item["case"] == "scanned_pdf")["httpStatus"] == 200
                and (
                    next(item for item in results if item["case"] == "scanned_pdf")["findingCount"] > 0
                    or next(item for item in results if item["case"] == "scanned_pdf")["limitedScan"]
                )
            ),
            "garbledFailsClosed": (
                next(item for item in results if item["case"] == "garbled")["httpStatus"] == 200
                and next(item for item in results if item["case"] == "garbled")["limitedScan"]
            ),
            "duplicateReplayedWithoutCharge": (
                next(item for item in results if item["case"] == "duplicate_replay")["httpStatus"] == 200
                and next(item for item in results if item["case"] == "duplicate_replay")["replayed"]
                and next(item for item in results if item["case"] == "duplicate_replay")["creditDelta"] == 0
            ),
        },
    }
    write_summary(output)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"Acceptance runner failed safely: {type(exc).__name__}: {exc}", file=sys.stderr)
        raise
